import streamlit as st
import pandas as pd
import numpy as np
import zipfile, pymupdf, re, os, tempfile
from docx import Document

df_data = {
    "Standard": [],
    "Matching Number": [],
    "Filename": [],
}

def get_standard_names():
    standard_names = os.listdir('standards/')

    return standard_names

def zip_outputs():
    with zipfile.ZipFile('outputs.zip', 'w') as zipf:
        for filename in os.listdir('outputs/'):
            file_path = 'outputs/' + filename
            zipf.write(file_path, compress_type=zipfile.ZIP_DEFLATED)

def clear_folder():
    files = os.listdir('outputs/')
    for file in files:
        os.remove(os.path.join('outputs/', file))

def get_match_data():
    df_data['Standard'] = []
    df_data['Matching Number'] = []
    df_data['Filename'] = []
    files = os.listdir('outputs/')
    for file in files:
        list_name, matching_number, filename = file.split('&')
        if matching_number != '0':
            df_data['Standard'].append(list_name)
            df_data['Matching Number'].append(matching_number)
            df_data['Filename'].append(filename)
    return pd.DataFrame(df_data)

class CASMatcher:
    def __init__(self):
        return

    def is_single_digit(self, string):
        return len(string) == 1 and string.isdigit()

    def find_numeric_hyphen_strings(self, string):
        pattern_1 = r'^\d+-\d+$'
        pattern_2 = r'^\d+-\d+-\d+$'
        if re.match(pattern_1, string) or re.match(pattern_2, string):
            return True
        return False

    def fit_excel(self, writer):
        for sheet_name in writer.sheets:
            worksheet = writer.sheets[sheet_name]
            for col in worksheet.columns:
                max_length = 0
                col_name = col[0].column_letter
                for cell in col:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)
                worksheet.column_dimensions[col_name].width = max_length

    def get_MDSReport_docx(self, report):
        doc = Document(report)
        data = []
        for table in doc.tables:
            for row in table.rows:
                temp = row.cells[0].text.strip(" |-,")
                if self.is_single_digit(temp):
                    data.append([cell.text for cell in row.cells])
        data = np.array(data)[:,0:3]
        data = np.char.strip(data.astype(str), " |-,")
        df = pd.DataFrame(data)
        df = df.rename(columns={df.columns[0]: 'Level', df.columns[1]: 'Substance Name', df.columns[2]: 'CAS Number'})
        return df

    def get_MDSReport_pdf(self, report):
        temp_file = tempfile.NamedTemporaryFile(delete=False)
        temp_file.write(report.read())
        temp_file.close()
        doc = pymupdf.Document(temp_file.name)
        dfs = []
        for page in doc:
            tabs = page.find_tables()
            if len(tabs.tables):
                for tab in tabs:
                    data = np.array(tab.extract())[:,0:3]
                    data = np.char.strip(data.astype(str), " |-,")
                    df = pd.DataFrame(data)
                    df = df[df.iloc[:, 0].apply(self.is_single_digit)]
                    df = df.rename(columns={df.columns[0]: 'Level', df.columns[1]: 'Substance Name', df.columns[2]: 'CAS Number'})
                    dfs.append(df)
        merged_df = pd.concat(dfs, ignore_index=True)
        return merged_df

    def preprocess_compareList(self, standard_list):
        df = pd.read_excel(standard_list)
        df = df.loc[:, ['CAS Number', 'Chemical Name']]
        df['CAS Number'] = df['CAS Number'].str.strip(" ,")
        df['CAS Number'] = df['CAS Number'].str.split(',')
        df = df.explode('CAS Number').reset_index(drop=True)
        df.to_excel('standards/' + 'processed_' + standard_list.name.split('/')[-1], index=False)
        return df

    def get_result(self, report, standard_name):
        file_type = report.name.split('.')[-1]
        compareList = pd.read_excel('standards/' + standard_name)
        if file_type == 'docx':
            data_MDS = self.get_MDSReport_docx(report)
        else:
            data_MDS = self.get_MDSReport_pdf(report)       
        result = pd.merge(
            data_MDS,
            compareList, on='CAS Number',
            how='left')
        count = result['Chemical Name'].notnull().sum()
        result['Chemical Name'] = result['Chemical Name'].fillna('---')
        level = 0
        yielded_rows = []
        for _, row in result[::-1].iterrows():
            currL = int(row['Level'])
            if row['Chemical Name'] != '---' or level > currL:
                yielded_rows.append(row)
                level = currL
        result = result[result.iloc[:, 2].apply(self.find_numeric_hyphen_strings)]
        result_Found = pd.DataFrame(yielded_rows[::-1])
        filename = report.name.split('/')[-1].split('.')[0] + '.xlsx'
        filepath = 'outputs/' + standard_name.split('_')[-1].split('.')[0] + '&' + str(count) + "&" + filename
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            result.to_excel(writer, sheet_name="Total List", index=False)
            if count:
                result_Found.to_excel(writer, sheet_name="Summary", index=False)
            self.fit_excel(writer)

def main():
    matcher = CASMatcher()
    st.set_page_config(page_title='CASMatcher', page_icon='https://www.johnsonelectric.com/pub/media/favicon/stores/1/johnson.jpg')
    st.title('CASMatcher Application')
    st.header('1. Upload the standard lists for comparison:', divider='rainbow')
    standard_lists = st.file_uploader(
        label='Upload Standard Lists',
        type=["xlsx"],
        accept_multiple_files=True,
        )
    if standard_lists is not None:
        for standard_list in standard_lists:
            matcher.preprocess_compareList(standard_list)
        standard_lists = None
    st.header('2. Choose the list for matching:', divider='rainbow')
    standard_names = get_standard_names()
    standard_name = st.selectbox(
        label='Select Standard List',
        options=standard_names,
        )
    if standard_name is not None:
        st.header('3. Upload the MAS reports:', divider='rainbow')
        MDSreports = st.file_uploader(
            label='Upload MAS Reports',
            type=["docx", "pdf"],
            accept_multiple_files=True,
            )
        _, _, body_col3, _, _ = st.columns(5)
        with body_col3:
            process = st.button("Process")
        progress_bar = st.progress(0, text="")
        if process and MDSreports is not None:
            if len(MDSreports) == 0:
                st.text("No reports are uploaded.")
            else:
                num = len(MDSreports)
                for i, report in enumerate(MDSreports):
                    matcher.get_result(report, standard_name)
                    progress_bar.progress(1*(i+1)/num, text="Operation in progress. Please wait.")
                progress_bar.progress(1., text="Operation Finished.")
                zip_outputs()
            MDSreports = None
        st.divider()
        end_col1, end_col2 = st.columns(2)
        with end_col1:
            st.subheader("Download the output files:")
            with open('outputs.zip', 'rb') as datazip:
                st.download_button(
                    label='Download ZIP',
                    data=datazip,
                    file_name="outputs.zip",
                    mime="application/octet-stream"
                    )
        with end_col2:
            st.subheader("Clear the output files:")
            st.button(label='Clear Outputs', on_click=clear_folder)
        st.subheader("List of files with matching CAS number:")
        st.dataframe(
            get_match_data(),
            hide_index=True,
        )
    else:
        st.caption('Please upload and select standard list for further actions.')
    
    with st.sidebar:
        st.header('	:books: Guideline')
        st.markdown('**1. Upload standard lists that you want to use for comparison.**')
        st.markdown('**2. Select the standard list for comparison.**')
        st.markdown('**3. Upload the MAS Reports before process.**')
        st.caption('*Please wait until all reports are uploaded for next step.')
        st.markdown('**4. Press the "Process" button.**')
        st.caption('*Wait until operation success.')
        st.markdown('**5. Download the output files via "Download ZIP" button.**')

if __name__ == '__main__':
    main()